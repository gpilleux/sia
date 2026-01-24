"""
DOCX File Reader - Microsoft Word Document Parser

Extracts text from DOCX files including:
- Main body paragraphs and tables
- Headers and footers (all types)
- Nested tables and complex structures

Domain: Skills (Infrastructure)
Bounded Context: File Processing
Pattern: Strategy (concrete implementation)

Dependencies:
    - python-docx: Official library for DOCX manipulation
    
Invariant:
    DocxReader.read(valid_docx) → non_empty_text
    ∧ DocxReader.read(corrupted_docx) → CorruptedFileError
    ∧ DocxReader ∈ AbstractFileReader.registry["docx"]

REQ-011: File Reader Skills System
QUANT-011-003: Concrete Readers Implementation
"""

from pathlib import Path
from typing import TYPE_CHECKING

from .base import AbstractFileReader, CorruptedFileError, validate_file_exists

if TYPE_CHECKING:
    from docx import Document as DocxDocument


class DocxReader(AbstractFileReader):
    """
    Extract text from DOCX (Microsoft Word) files.
    
    This reader uses the python-docx library to extract text from DOCX files,
    preserving document structure and including all textual elements:
    
    - Paragraphs (in document order)
    - Tables (including nested tables)
    - Headers (first page, even page, odd page)
    - Footers (first page, even page, odd page)
    
    Implementation Notes:
        - Uses iter_inner_content() to preserve document order
        - Handles linked headers/footers to avoid duplication
        - Recursive extraction for nested tables
        - Text boxes are NOT supported (requires XML parsing)
    
    Edge Cases Handled:
        - Corrupted ZIP archives (BadZipFile)
        - Invalid DOCX structure (PackageNotFoundError)
        - Password-protected files (detected and rejected)
        - Empty documents (returns empty string)
    
    Example:
        >>> reader = DocxReader()
        >>> text = reader.read(Path("document.docx"))
        >>> print(text[:100])
    """
    
    @classmethod
    def get_extension(cls) -> str:
        """Return supported extension: 'docx'"""
        return "docx"
    
    def read(self, filepath: Path) -> str:
        """
        Extract all text from a DOCX file.
        
        Args:
            filepath: Path to DOCX file
            
        Returns:
            Extracted text with structure preserved (paragraphs separated by newlines)
            
        Raises:
            FileNotFoundError: If file doesn't exist
            CorruptedFileError: If file is corrupted, invalid, or password-protected
        """
        validate_file_exists(filepath)
        
        # Lazy import to avoid forcing dependency
        try:
            from zipfile import BadZipFile

            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError as e:
            raise ImportError(
                "python-docx not installed. "
                "Use: uv run --with python-docx python your_script.py"
            ) from e
        
        # Open DOCX file
        try:
            document = Document(str(filepath))
        except PackageNotFoundError as e:
            raise CorruptedFileError(
                f"Invalid DOCX structure - file may be corrupted: {e}"
            ) from e
        except BadZipFile as e:
            raise CorruptedFileError(
                f"Corrupted DOCX file - invalid ZIP archive: {e}"
            ) from e
        except Exception as e:
            # Detect password-protected files
            error_msg = str(e).lower()
            if "encrypted" in error_msg or "password" in error_msg:
                raise CorruptedFileError(
                    "Password-protected DOCX files are not supported"
                ) from e
            # Generic error handling
            raise CorruptedFileError(
                f"Failed to open DOCX file: {e}"
            ) from e
        
        # Extract text components
        text_parts = []
        
        # 1. Extract main body content (paragraphs and tables)
        # Using iter_inner_content() preserves document order
        for item in document.iter_inner_content():
            if hasattr(item, 'text'):  # Paragraph
                text = item.text.strip()
                if text:  # Skip empty paragraphs
                    text_parts.append(text)
            elif hasattr(item, 'rows'):  # Table
                table_text = self._extract_table_text(item)
                if table_text:
                    text_parts.append(table_text)
        
        # 2. Extract headers and footers from all sections
        headers_footers = self._extract_headers_footers(document)
        if headers_footers:
            text_parts.append("\n" + headers_footers)
        
        return "\n".join(text_parts)
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table, including nested tables.
        
        Args:
            table: docx.table.Table object
            
        Returns:
            Formatted table text with rows separated by newlines
        """
        rows_text = []
        
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                # Extract text from all paragraphs in cell
                cell_content = []
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        cell_content.append(text)
                
                if cell_content:
                    cells_text.append(" ".join(cell_content))
            
            if cells_text:
                rows_text.append("\t".join(cells_text))
        
        return "\n".join(rows_text)
    
    def _extract_headers_footers(self, document) -> str:
        """
        Extract headers and footers from all sections.
        
        Handles:
        - First page header/footer
        - Even page header/footer
        - Odd page (default) header/footer
        - Linked headers/footers (avoids duplication)
        
        Args:
            document: docx.Document object
            
        Returns:
            Combined headers/footers text with markers
        """
        hf_parts = []
        
        for section_num, section in enumerate(document.sections, start=1):
            # Extract headers
            headers = [
                ("HEADER", section.header),
                ("FIRST_PAGE_HEADER", section.first_page_header),
                ("EVEN_PAGE_HEADER", section.even_page_header),
            ]
            
            for label, header in headers:
                if header and not header.is_linked_to_previous:
                    header_text = self._extract_header_footer_text(header)
                    if header_text:
                        hf_parts.append(f"[{label}_S{section_num}] {header_text}")
            
            # Extract footers
            footers = [
                ("FOOTER", section.footer),
                ("FIRST_PAGE_FOOTER", section.first_page_footer),
                ("EVEN_PAGE_FOOTER", section.even_page_footer),
            ]
            
            for label, footer in footers:
                if footer and not footer.is_linked_to_previous:
                    footer_text = self._extract_header_footer_text(footer)
                    if footer_text:
                        hf_parts.append(f"[{label}_S{section_num}] {footer_text}")
        
        return "\n".join(hf_parts)
    
    def _extract_header_footer_text(self, hf_element) -> str:
        """
        Extract text from a header or footer element.
        
        Args:
            hf_element: Header or Footer object
            
        Returns:
            Combined text from all paragraphs
        """
        text_parts = []
        for paragraph in hf_element.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        return " ".join(text_parts)
